import re
import uuid
from dataclasses import dataclass, field

from docx import Document

PLACEHOLDER_RE = re.compile(r"\{([^{}]+)\}")

# Стандартные коды блоков по §2.4 SPEC — используются как база, если из docx
# не удаётся однозначно определить код блока по заголовку.
DEFAULT_BLOCK_CODES = {
    "цел": "goals",
    "задач": "goals",
    "периметр": "scope",
    "объект": "scope",
    "срок": "terms",
    "содержание работ": "work_content",
    "этап": "work_content",
    "услови": "conditions",
    "документ": "documentation",
    "контрол": "quality_control",
    "качеств": "quality_control",
    "подпис": "signatures",
}


def _guess_block_code(heading_text: str, index: int) -> str:
    lowered = heading_text.lower()
    for key, code in DEFAULT_BLOCK_CODES.items():
        if key in lowered:
            return code
    return f"block_{index}"


@dataclass
class ParsedBlock:
    code: str
    name: str
    order: int
    fields: list[dict] = field(default_factory=list)
    is_stages_block: bool = False


@dataclass
class ParsedTemplate:
    blocks: list[ParsedBlock]
    stages: list[dict]
    raw_docx: bytes


def parse_docx_template(file_bytes: bytes) -> ParsedTemplate:
    import io

    document = Document(io.BytesIO(file_bytes))

    blocks: list[ParsedBlock] = []
    current_block: ParsedBlock | None = None
    order = 0

    for para in document.paragraphs:
        style_name = (para.style.name or "").lower() if para.style else ""
        text = para.text.strip()
        if not text:
            continue
        if "heading" in style_name or "заголовок" in style_name:
            order += 1
            code = _guess_block_code(text, order)
            current_block = ParsedBlock(
                code=code,
                name=text,
                order=order,
                is_stages_block=(code == "work_content"),
            )
            blocks.append(current_block)
            continue

        placeholders = PLACEHOLDER_RE.findall(text)
        if current_block is not None:
            for placeholder in placeholders:
                key = re.sub(r"[^a-zA-Zа-яА-Я0-9_]+", "_", placeholder).strip("_").lower() or "field"
                current_block.fields.append(
                    {
                        "key": key,
                        "type": "text",
                        "label": placeholder,
                        "placeholder": f"{{{placeholder}}}",
                        "required": False,
                    }
                )
            if not placeholders and not any(f["label"] == text for f in current_block.fields):
                current_block.fields.append(
                    {"key": f"text_{len(current_block.fields)}", "type": "text", "label": text[:80], "required": False}
                )

    stages: list[dict] = []
    for table in document.tables:
        header_cells = [c.text.strip().lower() for c in table.rows[0].cells] if table.rows else []
        if any("требован" in h for h in header_cells) or any("результат" in h for h in header_cells):
            for i, row in enumerate(table.rows[1:], start=1):
                cells = [c.text.strip() for c in row.cells]
                if not any(cells):
                    continue
                stage_name = cells[0] if cells else f"Этап {i}"
                requirements = cells[1] if len(cells) > 1 else None
                results = cells[2] if len(cells) > 2 else None
                stages.append(
                    {
                        "stage_order": i,
                        "stage_name": stage_name or f"Этап {i}",
                        "default_requirements": requirements,
                        "default_results": results,
                    }
                )

    if not blocks:
        blocks = _default_blocks()

    return ParsedTemplate(blocks=blocks, stages=stages, raw_docx=file_bytes)


def _default_blocks() -> list[ParsedBlock]:
    return [
        ParsedBlock(code="goals", name="Цели и задачи работ", order=1, fields=[
            {"key": "goal_text", "type": "text", "label": "Цель", "required": True},
            {"key": "tasks", "type": "list", "label": "Задачи", "required": True},
        ]),
        ParsedBlock(code="scope", name="Периметр работ", order=2, fields=[
            {"key": "location", "type": "text", "label": "Место оказания", "required": False},
            {"key": "field_name", "type": "text", "label": "Наименование месторождения", "required": True},
        ]),
        ParsedBlock(code="terms", name="Сроки выполнения работ", order=3, fields=[
            {"key": "date_start", "type": "date", "label": "Начало", "required": True},
            {"key": "date_end", "type": "date", "label": "Окончание", "required": True},
        ]),
        ParsedBlock(code="work_content", name="Содержание работ", order=4, is_stages_block=True),
        ParsedBlock(code="conditions", name="Условия выполнения работы", order=5, fields=[
            {"key": "source_data", "type": "text", "label": "Исходная информация от Заказчика", "required": False},
            {"key": "software", "type": "text", "label": "Программное обеспечение", "required": False},
        ]),
        ParsedBlock(code="documentation", name="Требования к документации", order=6, fields=[
            {"key": "report_formats", "type": "text", "label": "Форматы отчётов", "required": False},
        ]),
        ParsedBlock(code="quality_control", name="Контроль качества", order=7, fields=[
            {"key": "acceptance", "type": "text", "label": "Условия приёмки", "required": False},
        ]),
        ParsedBlock(code="signatures", name="Подписи сторон", order=8, fields=[
            {"key": "customer_signee", "type": "text", "label": "Подписант Заказчика", "required": True},
            {"key": "contractor_signee", "type": "text", "label": "Подписант Исполнителя", "required": True},
        ]),
    ]


def build_blocks_schema(blocks: list[ParsedBlock]) -> dict:
    return {
        "blocks": [
            {
                "code": b.code,
                "name": b.name,
                "order": b.order,
                **({"is_stages_block": True} if b.is_stages_block else {"fields": b.fields}),
            }
            for b in blocks
        ]
    }


def new_template_docx_key(template_id: uuid.UUID) -> str:
    return f"templates/{template_id}.docx"
