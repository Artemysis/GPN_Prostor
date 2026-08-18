import io
import uuid

from docx import Document
from docx.shared import Pt

from app.core.config import get_settings
from app.db.models import Request, RequestTz, RequestTzStage, TzTemplate
from app.services.llm_client import DeepSeekLLMClient, get_llm_client
from app.services.minio_client import MinioService

settings = get_settings()


def render_tz_docx(request: Request, template: TzTemplate, tz: RequestTz, stages: list[RequestTzStage]) -> bytes:
    doc = Document()
    doc.add_heading(f"Техническое задание: {template.name}", level=0)

    meta = doc.add_paragraph()
    meta.add_run(f"Заявка №{request.number or request.id}\n").bold = True
    meta.add_run(f"Наименование: {request.title or '-'}\n")
    meta.add_run(f"Заказчик/подрядчик: {request.company_id or '-'}\n")
    meta.add_run(f"Сроки: {request.date_start or '-'} — {request.date_end or '-'}\n")

    blocks = sorted(template.blocks_schema.get("blocks", []), key=lambda b: b.get("order", 0))
    for block in blocks:
        code = block["code"]
        doc.add_heading(block.get("name", code), level=1)
        content = tz.payload.get(code, {}) or {}
        if block.get("is_stages_block"):
            table = doc.add_table(rows=1, cols=4)
            table.style = "Light Grid Accent 1"
            hdr = table.rows[0].cells
            hdr[0].text, hdr[1].text, hdr[2].text, hdr[3].text = "Этап", "Описание", "Требования", "Ожидаемые результаты"
            for stage in sorted(stages, key=lambda s: s.stage_order):
                row = table.add_row().cells
                row[0].text = stage.stage_name
                row[1].text = stage.description or ""
                row[2].text = stage.requirements or ""
                row[3].text = stage.expected_results or ""
        else:
            for field in block.get("fields", []):
                key = field["key"]
                label = field.get("label", key)
                value = content.get(key)
                if isinstance(value, list):
                    value = "; ".join(str(v) for v in value)
                p = doc.add_paragraph()
                p.add_run(f"{label}: ").bold = True
                p.add_run(str(value) if value else "—")

    doc.add_heading("Оценочная стоимость", level=1)
    cost_p = doc.add_paragraph()
    cost_p.add_run("Стоимость: ").bold = True
    cost_p.add_run(f"{request.cost_total or '-'} {request.currency}")

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


ANALYTICAL_REPORT_SYSTEM_PROMPT = (
    "Ты — аналитик платформы ПРОСТОР. Сформируй аналитический отчёт по ТЗ и результатам анализа. "
    "Разделы: сводка, качество ТЗ, риски, рекомендации, сравнение с типовым, прогноз сроков/стоимости. "
    "Стиль — деловой, для руководителя. Ответ — обычный текст (не JSON)."
)


async def generate_analytical_report_text(
    request: Request,
    tz: RequestTz,
    analysis: dict,
    llm: DeepSeekLLMClient | None = None,
) -> str:
    llm = llm or get_llm_client()
    model = settings.llm_reasoner_model if settings.llm_use_reasoner_for_report else settings.llm_model
    user_prompt = f"Заявка: {request.title}. ТЗ payload: {tz.payload}. Результаты анализа: {analysis}."
    text = await llm.chat_text(ANALYTICAL_REPORT_SYSTEM_PROMPT, user_prompt, model=model)
    if not text:
        text = _fallback_report_text(request, analysis)
    return text


def _fallback_report_text(request: Request, analysis: dict) -> str:
    lines = [
        f"Аналитический отчёт по заявке {request.number or request.id}",
        "",
        f"Готовность ТЗ: {analysis.get('completeness_pct', 0)}%",
        "",
        "Риски:",
    ]
    for risk in analysis.get("risks", []):
        lines.append(f"- [{risk.get('severity')}] {risk.get('title')}: {risk.get('description')}")
    lines.append("")
    lines.append("Рекомендации:")
    for rec in analysis.get("recommendations", []):
        lines.append(f"- {rec.get('title')}: {rec.get('description')}")
    return "\n".join(lines)


def render_report_docx(title: str, text: str) -> bytes:
    doc = Document()
    doc.add_heading(title, level=0)
    for paragraph in text.split("\n"):
        p = doc.add_paragraph(paragraph)
        p.style.font.size = Pt(11)
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def render_text_pdf(title: str, text: str) -> bytes:
    from weasyprint import HTML

    escaped = text.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;").replace("\n", "<br/>")
    html = f"<html><head><meta charset='utf-8'></head><body><h1>{title}</h1><p>{escaped}</p></body></html>"
    return HTML(string=html).write_pdf()


def tz_payload_as_text(request: Request, template: TzTemplate, tz: RequestTz, stages: list[RequestTzStage]) -> str:
    lines = [f"Техническое задание: {template.name}", f"Заявка №{request.number or request.id}", ""]
    blocks = sorted(template.blocks_schema.get("blocks", []), key=lambda b: b.get("order", 0))
    for block in blocks:
        code = block["code"]
        lines.append(block.get("name", code))
        content = tz.payload.get(code, {}) or {}
        if block.get("is_stages_block"):
            for stage in sorted(stages, key=lambda s: s.stage_order):
                lines.append(f"  - {stage.stage_name}: {stage.requirements or ''} / {stage.expected_results or ''}")
        else:
            for field in block.get("fields", []):
                lines.append(f"  {field.get('label', field['key'])}: {content.get(field['key'], '—')}")
        lines.append("")
    return "\n".join(lines)


def export_key(request_id: uuid.UUID, document_id: uuid.UUID, ext: str = "docx") -> str:
    return f"exports/{request_id}/{document_id}.{ext}"


_EXPORT_CONTENT_TYPES = {
    "docx": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    "xlsx": "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
    "pdf": "application/pdf",
}


def upload_export(minio: MinioService, request_id: uuid.UUID, document_id: uuid.UUID, data: bytes, ext: str = "docx") -> tuple[str, str]:
    key = export_key(request_id, document_id, ext)
    content_type = _EXPORT_CONTENT_TYPES.get(ext, "application/octet-stream")
    minio.upload_bytes(settings.minio_bucket_exports, key, data, content_type)
    return settings.minio_bucket_exports, key
