"""Генерация комплекта документов заявки (Наряд-заказ, Приложение 1 ТЗ,
Приложение № 2.1 Форма ТЗ, Приложение 2 КП, Приложение 3 РС).

Структура документов повторяет шаблоны из seed/package (см. также
`app.services.seed.seed_package_templates`, которая загружает исходные файлы
в MinIO при старте backend). Итоговые документы не обязаны быть побайтовой
копией шаблонов, но должны сохранять их структуру (разделы, таблицы, подписи).
"""

import io
import uuid
from datetime import date

import openpyxl
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from openpyxl.styles import Alignment, Font

from app.db.models import (
    Company,
    Contract,
    Product,
    ProductRate,
    Request,
    RequestTz,
    RequestTzStage,
    TzTemplate,
)

VAT_RATE = 0.22

ROLE_FALLBACK = ["Ведущий специалист", "Специалист", "Инженер-аналитик", "Технический эксперт"]


def _fmt_date(value: date | None) -> str:
    return value.strftime("%d.%m.%Y") if value else "не указана"


def _fmt_money(value: float) -> str:
    return f"{value:,.2f}".replace(",", " ").replace(".", ",")


def _block(tz: RequestTz, code: str) -> dict:
    return (tz.payload or {}).get(code, {}) or {}


def _sorted_stages(stages: list[RequestTzStage]) -> list[RequestTzStage]:
    return sorted(stages, key=lambda s: s.stage_order)


def _stage_costs(stages: list[RequestTzStage], cost_total: float) -> list[float]:
    """Распределяет общую стоимость заявки по этапам пропорционально длительности."""
    if not stages:
        return []
    durations = []
    for stage in stages:
        if stage.stage_start_date and stage.stage_end_date:
            durations.append(max((stage.stage_end_date - stage.stage_start_date).days, 1))
        else:
            durations.append(1)
    total_duration = sum(durations) or len(stages)
    costs = [cost_total * d / total_duration for d in durations]
    # компенсируем ошибку округления в последнем этапе
    if costs:
        costs[-1] += cost_total - sum(costs)
    return costs


def _bold_run(paragraph, text: str) -> None:
    paragraph.add_run(text).bold = True


# ---------------------------------------------------------------------------
# Наряд-заказ (docx)
# ---------------------------------------------------------------------------


def render_naryad_zakaz_docx(
    request: Request,
    company: Company | None,
    contract: Contract | None,
    product: Product | None,
    tz: RequestTz | None,
) -> bytes:
    doc = Document()
    title = doc.add_heading("Наряд-заказ", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run(f"к Договору № {contract.contract_number if contract else '—'}")

    p = doc.add_paragraph()
    p.add_run(f"г. Санкт-Петербург\t\t\t\t{_fmt_date(date.today())}")

    doc.add_paragraph(
        "Заказчик и Исполнитель, именуемые в дальнейшем Стороны, составили настоящий наряд-заказ "
        "о нижеследующем на основании заявки, оформленной в системе ПРОСТОР."
    )

    doc.add_heading("1. Основные условия сделки", level=1)
    table = doc.add_table(rows=0, cols=2)
    table.style = "Light Grid Accent 1"
    rows = [
        ("Наименование сделки", request.title or "—"),
        ("Исполнитель", company.name if company else "—"),
        ("Номер договора", contract.contract_number if contract else "—"),
        ("Продукт / услуга", product.product_name if product else "—"),
        ("Стоимость сделки", f"{_fmt_money(float(request.cost_total or 0))} {request.currency}"),
        ("Срок выполнения работ", f"{_fmt_date(request.date_start)} — {_fmt_date(request.date_end)}"),
    ]
    for label, value in rows:
        row = table.add_row().cells
        row[0].text = label
        row[1].text = str(value)

    doc.add_heading("2. Прочие условия", level=1)
    goals = _block(tz, "goals") if tz else {}
    conditions = _block(tz, "conditions") if tz else {}
    table2 = doc.add_table(rows=0, cols=2)
    table2.style = "Light Grid Accent 1"
    rows2 = [
        ("Предмет и цель работ", goals.get("goal_text", "—")),
        ("Исходные данные от Заказчика", conditions.get("source_data", "—")),
        ("Программное обеспечение", conditions.get("software", "—")),
        ("Количество экземпляров наряда-заказа", "2 (два) экземпляра, по одному для каждой из Сторон"),
    ]
    for label, value in rows2:
        row = table2.add_row().cells
        row[0].text = label
        row[1].text = str(value or "—")

    doc.add_paragraph(
        "3. Настоящий наряд-заказ вступает в силу с даты подписания обеими Сторонами и действует "
        f"до {_fmt_date(request.date_end)}."
    )

    doc.add_heading("Подписи сторон", level=1)
    sig = doc.add_table(rows=1, cols=2)
    sig.style = "Light Grid Accent 1"
    cells = sig.rows[0].cells
    cells[0].text = "От Заказчика:\nдолжность\n\n________________ /___________/"
    cells[1].text = f"От Исполнителя:\nдолжность\n\n________________ /{company.name if company else '___________'}/"

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


# ---------------------------------------------------------------------------
# Приложение 1. ТЗ (docx) — полная версия, повторяет структуру шаблона
# ---------------------------------------------------------------------------


def render_tz_appendix1_docx(
    request: Request, template: TzTemplate, tz: RequestTz, stages: list[RequestTzStage]
) -> bytes:
    doc = Document()
    doc.add_heading("Приложение №1", level=0)
    doc.add_paragraph(f"к заявке № {request.number or request.id} от {_fmt_date(date.today())}")
    doc.add_paragraph(f"Техническое задание: {template.name}")

    blocks = sorted(template.blocks_schema.get("blocks", []), key=lambda b: b.get("order", 0))
    for block in blocks:
        code = block["code"]
        doc.add_heading(block.get("name", code), level=1)
        content = _block(tz, code)
        if block.get("is_stages_block"):
            table = doc.add_table(rows=1, cols=5)
            table.style = "Light Grid Accent 1"
            hdr = table.rows[0].cells
            hdr[0].text, hdr[1].text, hdr[2].text, hdr[3].text, hdr[4].text = (
                "№ этапа",
                "Наименование этапа",
                "Описание работ",
                "Требования к выполнению",
                "Ожидаемые результаты",
            )
            for stage in _sorted_stages(stages):
                row = table.add_row().cells
                row[0].text = str(stage.stage_order)
                row[1].text = stage.stage_name
                row[2].text = stage.description or "—"
                row[3].text = stage.requirements or "—"
                row[4].text = stage.expected_results or "—"
        else:
            for field in block.get("fields", []):
                key = field["key"]
                label = field.get("label", key)
                value = content.get(key)
                if isinstance(value, list):
                    value = "; ".join(str(v) for v in value)
                p = doc.add_paragraph()
                _bold_run(p, f"{label}: ")
                p.add_run(str(value) if value else "—")

    doc.add_heading("Подписи сторон", level=1)
    sig = doc.add_table(rows=1, cols=2)
    sig.style = "Light Grid Accent 1"
    cells = sig.rows[0].cells
    signatures = _block(tz, "signatures")
    cells[0].text = f"От Заказчика:\n{signatures.get('customer_signee') or '___________'}\n\n________________"
    cells[1].text = f"От Исполнителя:\n{signatures.get('contractor_signee') or '___________'}\n\n________________"

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


# ---------------------------------------------------------------------------
# Приложение № 2.1. Форма Технического задания (docx) — краткая форма
# ---------------------------------------------------------------------------


def render_tz_form_2_1_docx(
    request: Request, template: TzTemplate, tz: RequestTz, stages: list[RequestTzStage]
) -> bytes:
    doc = Document()
    doc.add_heading("Приложение № 2.1", level=0)
    doc.add_paragraph(f"к наряду-заказу № {request.number or request.id} от {_fmt_date(date.today())}")
    doc.add_heading("Форма Технического задания", level=1)

    doc.add_heading("Раздел 1. Описание работ / услуг", level=2)
    goals = _block(tz, "goals")
    scope = _block(tz, "scope")
    p = doc.add_paragraph()
    _bold_run(p, "Наименование работ/услуг: ")
    p.add_run(template.name)
    p = doc.add_paragraph()
    _bold_run(p, "Цель работ: ")
    p.add_run(goals.get("goal_text") or "—")
    tasks = goals.get("tasks") or []
    if tasks:
        doc.add_paragraph("Задачи:")
        for task in tasks:
            doc.add_paragraph(str(task), style="List Bullet")
    p = doc.add_paragraph()
    _bold_run(p, "Объект работ: ")
    p.add_run(scope.get("field_name") or "—")
    p = doc.add_paragraph()
    _bold_run(p, "Место выполнения работ: ")
    p.add_run(scope.get("location") or "—")

    doc.add_heading("Раздел 2. Сроки выполнения работ", level=2)
    doc.add_paragraph(f"Начало: {_fmt_date(request.date_start)}")
    doc.add_paragraph(f"Окончание: {_fmt_date(request.date_end)}")

    doc.add_heading("Раздел 3. Требования к результату работ", level=2)
    table = doc.add_table(rows=1, cols=3)
    table.style = "Light Grid Accent 1"
    hdr = table.rows[0].cells
    hdr[0].text, hdr[1].text, hdr[2].text = "№ этапа", "Наименование этапа", "Ожидаемый результат"
    for stage in _sorted_stages(stages):
        row = table.add_row().cells
        row[0].text = str(stage.stage_order)
        row[1].text = stage.stage_name
        row[2].text = stage.expected_results or "—"

    doc.add_heading("Раздел 4. Требования к отчётности и контролю качества", level=2)
    documentation = _block(tz, "documentation")
    quality = _block(tz, "quality_control")
    doc.add_paragraph(f"Форматы отчётов: {documentation.get('report_formats') or '—'}")
    doc.add_paragraph(f"Условия приёмки: {quality.get('acceptance') or '—'}")

    doc.add_heading("Подписи сторон", level=1)
    sig = doc.add_table(rows=1, cols=2)
    sig.style = "Light Grid Accent 1"
    cells = sig.rows[0].cells
    cells[0].text = "Заказчик:\n\n________________ /___/"
    cells[1].text = "Исполнитель:\n\n________________ /___/"

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


# ---------------------------------------------------------------------------
# Приложение 2. КП (xlsx)
# ---------------------------------------------------------------------------


def _kp_header(ws, request: Request) -> None:
    ws["F1"] = "Приложение 2"
    ws["F2"] = f"к наряду-заказу № {request.number or request.id} от {_fmt_date(date.today())}"
    ws["F3"] = f"к договору № {request.contract_id or '—'} от ___________"
    ws["A5"] = "Коммерческое предложение"
    ws["A5"].font = Font(bold=True, size=12)
    ws["A6"] = "на выполнение работ по теме"
    ws["A7"] = request.title or "—"


def render_kp_xlsx(request: Request, stages: list[RequestTzStage], company: Company | None) -> bytes:
    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = "КП"
    _kp_header(ws, request)

    header_row = 9
    headers = [
        "№ п/п",
        "Наименование этапа",
        "Состав исполнителей / квалификация",
        "Начало",
        "Окончание",
        "Стоимость этапа, руб. (без НДС)",
    ]
    for col, text in enumerate(headers, start=1):
        cell = ws.cell(row=header_row, column=col, value=text)
        cell.font = Font(bold=True)
        cell.alignment = Alignment(wrap_text=True, vertical="center")

    cost_total = float(request.cost_total or 0)
    ordered = _sorted_stages(stages)
    costs = _stage_costs(ordered, cost_total)

    row = header_row + 1
    for stage, cost in zip(ordered, costs, strict=True):
        ws.cell(row=row, column=1, value=stage.stage_order)
        ws.cell(row=row, column=2, value=stage.stage_name)
        ws.cell(row=row, column=3, value="Профильные специалисты Исполнителя")
        ws.cell(row=row, column=4, value=_fmt_date(stage.stage_start_date))
        ws.cell(row=row, column=5, value=_fmt_date(stage.stage_end_date))
        ws.cell(row=row, column=6, value=round(cost, 2))
        row += 1

    if not ordered:
        ws.cell(row=row, column=2, value="Этапы не заданы — заполните конструктор ТЗ")
        row += 1

    vat = cost_total * VAT_RATE
    row += 1
    ws.cell(row=row, column=2, value="Всего без НДС, руб.")
    ws.cell(row=row, column=6, value=round(cost_total, 2))
    row += 1
    ws.cell(row=row, column=2, value="Всего НДС (22%), руб.")
    ws.cell(row=row, column=6, value=round(vat, 2))
    row += 1
    ws.cell(row=row, column=2, value="Итого с НДС, руб.")
    ws.cell(row=row, column=6, value=round(cost_total + vat, 2))

    row += 3
    ws.cell(row=row, column=2, value="Заказчик")
    ws.cell(row=row, column=4, value="Исполнитель")
    row += 2
    ws.cell(row=row, column=2, value="_____________________ /___/")
    ws.cell(row=row, column=4, value=f"________________ /{company.name if company else '___'}/")

    for col, width in zip("ABCDEF", [6, 40, 28, 14, 14, 22], strict=True):
        ws.column_dimensions[col].width = width

    buf = io.BytesIO()
    wb.save(buf)
    return buf.getvalue()


# ---------------------------------------------------------------------------
# Приложение 3. РС (xlsx)
# ---------------------------------------------------------------------------


def render_rs_xlsx(request: Request, stages: list[RequestTzStage], rates: list[ProductRate]) -> bytes:
    wb = openpyxl.Workbook()
    ws1 = wb.active
    ws1.title = "Лист1"

    ws1["F1"] = "Приложение № 3"
    ws1["F2"] = f"к наряду-заказу № {request.number or request.id}"
    ws1["A3"] = "Расчёт стоимости работ"
    ws1["A3"].font = Font(bold=True, size=12)

    ws1["A5"] = "1. Перечень выполняемых работ:"
    ws1["A5"].font = Font(bold=True)
    ws1["A6"] = "№ пп"
    ws1["B6"] = "Наименование этапа"
    for cell in ("A6", "B6"):
        ws1[cell].font = Font(bold=True)

    cost_total = float(request.cost_total or 0)
    ordered = _sorted_stages(stages)
    row = 7
    for stage in ordered:
        ws1.cell(row=row, column=1, value=stage.stage_order)
        ws1.cell(row=row, column=2, value=stage.stage_name)
        row += 1
    if not ordered:
        ws1.cell(row=row, column=2, value="Этапы не заданы — заполните конструктор ТЗ")
        row += 1

    vat = cost_total * VAT_RATE
    row += 1
    ws1.cell(row=row, column=2, value="Всего без НДС, руб.")
    ws1.cell(row=row, column=3, value=round(cost_total, 2))
    row += 1
    ws1.cell(row=row, column=2, value="Всего НДС (22%), руб.")
    ws1.cell(row=row, column=3, value=round(vat, 2))
    row += 1
    ws1.cell(row=row, column=2, value="Итого с НДС, руб.")
    ws1.cell(row=row, column=3, value=round(cost_total + vat, 2))

    row += 2
    ws1.cell(row=row, column=1, value="2. Расчёт стоимости (детализация — см. лист «Лист2»):")
    ws1.cell(row=row, column=1).font = Font(bold=True)

    ws1.column_dimensions["A"].width = 8
    ws1.column_dimensions["B"].width = 45
    ws1.column_dimensions["C"].width = 20

    ws2 = wb.create_sheet("Лист2")
    ws2["A1"] = "2. Расчёт стоимости (детализация по категориям исполнителей):"
    ws2["A1"].font = Font(bold=True)
    headers = [
        "№ п/п",
        "Категория",
        "Ставка в час, руб. без НДС",
        "Трудозатраты, FTE",
        "Человеко-дни",
        "Стоимость, руб.",
    ]
    for col, text in enumerate(headers, start=1):
        cell = ws2.cell(row=2, column=col, value=text)
        cell.font = Font(bold=True)
        cell.alignment = Alignment(wrap_text=True)

    role_names = [r.price_name for r in rates] or ROLE_FALLBACK
    role_names = role_names[:6] or ROLE_FALLBACK
    n = len(role_names)
    share = cost_total / n if n else 0
    hourly_rate = 1000.0

    row = 3
    running_cost = 0.0
    for idx, name in enumerate(role_names, start=1):
        cost_i = share
        if idx == n:
            cost_i = cost_total - running_cost
        running_cost += cost_i
        human_days = round(cost_i / hourly_rate / 8, 2) if hourly_rate else 0
        fte = round(human_days / 249, 2) if human_days else 0
        ws2.cell(row=row, column=1, value=idx)
        ws2.cell(row=row, column=2, value=name)
        ws2.cell(row=row, column=3, value=hourly_rate)
        ws2.cell(row=row, column=4, value=fte)
        ws2.cell(row=row, column=5, value=human_days)
        ws2.cell(row=row, column=6, value=round(cost_i, 2))
        row += 1

    row += 1
    ws2.cell(row=row, column=2, value="Всего без НДС, руб.")
    ws2.cell(row=row, column=6, value=round(cost_total, 2))
    row += 1
    ws2.cell(row=row, column=2, value="Всего НДС (22%), руб.")
    ws2.cell(row=row, column=6, value=round(vat, 2))
    row += 1
    ws2.cell(row=row, column=2, value="Итого с НДС, руб.")
    ws2.cell(row=row, column=6, value=round(cost_total + vat, 2))

    row += 2
    ws2.cell(row=row, column=2, value="Заказчик")
    ws2.cell(row=row, column=4, value="Исполнитель")
    row += 2
    ws2.cell(row=row, column=2, value="_______________/___/")
    ws2.cell(row=row, column=4, value="_______________/___/")

    for col, width in zip("ABCDEF", [6, 32, 16, 14, 14, 16], strict=True):
        ws2.column_dimensions[col].width = width

    buf = io.BytesIO()
    wb.save(buf)
    return buf.getvalue()


_DOCX_MIME = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
_XLSX_MIME = "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"

PACKAGE_DOCS = [
    # (kind, filename_prefix, mime_type, ext)
    ("naryad_zakaz", "Наряд-заказ", _DOCX_MIME, "docx"),
    ("tz_appendix1", "Приложение_1_ТЗ", _DOCX_MIME, "docx"),
    ("tz_form_2_1", "Приложение_2.1_Форма_ТЗ", _DOCX_MIME, "docx"),
    ("kp", "Приложение_2_КП", _XLSX_MIME, "xlsx"),
    ("rs", "Приложение_3_РС", _XLSX_MIME, "xlsx"),
]


def package_export_key(request_id: uuid.UUID, document_id: uuid.UUID, ext: str) -> str:
    return f"exports/{request_id}/{document_id}.{ext}"
